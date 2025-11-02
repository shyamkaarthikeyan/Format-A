#!/usr/bin/env python3
"""
Test script to verify Python serverless endpoints work correctly.
This simulates HTTP requests to the serverless functions.
"""

import sys
import json
import os
import base64
from io import BytesIO, StringIO
from unittest.mock import Mock

# Add the api directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'api'))

def create_mock_handler(module_name, method='POST', path='/api/test', body=None, headers=None):
    """Create a mock HTTP handler for testing."""
    
    # Import the module dynamically
    import importlib.util
    spec = importlib.util.spec_from_file_location(module_name, f"api/{module_name}.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    
    # Create mock handler
    class MockHandler(module.handler):
        def __init__(self):
            self.response_status = None
            self.response_headers = {}
            self.response_body = BytesIO()
            self.path = path
            self.headers = headers or {}
            self.wfile = self.response_body
            self.request_body = body or b''
        
        def send_response(self, status):
            self.response_status = status
        
        def send_header(self, name, value):
            self.response_headers[name] = value
        
        def end_headers(self):
            pass
        
        def rfile_read(self, length):
            return self.request_body[:length]
        
        # Override rfile to use our mock
        @property
        def rfile(self):
            mock_rfile = Mock()
            mock_rfile.read = self.rfile_read
            return mock_rfile
    
    return MockHandler()

def test_health_endpoint():
    """Test the health check endpoint."""
    print("Testing health check endpoint...")
    
    try:
        # Test GET request
        handler = create_mock_handler('health-python', 'GET', '/api/health-python')
        handler.do_GET()
        
        # Check response
        assert handler.response_status == 200 or handler.response_status == 206
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        
        assert response_json['success'] == True
        assert 'dependencies' in response_json
        assert 'capabilities' in response_json
        
        print(f"✓ Health check GET: Status {handler.response_status}")
        
        # Test POST request with detailed tests
        test_request = {
            'test_pdf_generation': True,
            'test_docx_conversion': True
        }
        
        handler = create_mock_handler(
            'health-python', 
            'POST', 
            '/api/health-python',
            json.dumps(test_request).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(test_request)))}
        )
        handler.do_POST()
        
        # Check response
        assert handler.response_status == 200 or handler.response_status == 206
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        
        assert response_json['success'] == True
        assert 'tests' in response_json
        
        print(f"✓ Health check POST: Status {handler.response_status}")
        print(f"  - PDF generation test: {response_json['tests']['pdf_generation']['status']}")
        print(f"  - DOCX conversion test: {response_json['tests']['docx_conversion']['status']}")
        
        return True
        
    except Exception as e:
        print(f"✗ Health check endpoint test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_pdf_generation_endpoint():
    """Test the PDF generation endpoint."""
    print("\nTesting PDF generation endpoint...")
    
    try:
        # Create a test document request
        test_document = {
            'title': 'Test IEEE Document',
            'authors': [
                {
                    'name': 'Test Author',
                    'department': 'Computer Science',
                    'organization': 'Test University',
                    'city': 'Test City',
                    'state': 'Test State'
                }
            ],
            'abstract': 'This is a test abstract for the IEEE document generation test.',
            'keywords': 'test, IEEE, document, generation',
            'sections': [
                {
                    'title': 'Introduction',
                    'contentBlocks': [
                        {
                            'type': 'text',
                            'content': 'This is the introduction section of the test document.'
                        }
                    ]
                },
                {
                    'title': 'Methodology',
                    'contentBlocks': [
                        {
                            'type': 'text',
                            'content': 'This section describes the methodology used in the test.'
                        }
                    ]
                }
            ]
        }
        
        # Test preview mode
        handler = create_mock_handler(
            'generate-pdf',
            'POST',
            '/api/generate-pdf?preview=true',
            json.dumps(test_document).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(test_document)))}
        )
        handler.do_POST()
        
        # Check response
        assert handler.response_status == 200
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        
        assert response_json['success'] == True
        assert 'data' in response_json
        assert response_json['format'] == 'docx'
        
        # Verify the base64 data can be decoded
        docx_data = base64.b64decode(response_json['data'])
        assert len(docx_data) > 0
        
        print(f"✓ PDF generation (preview): Status {handler.response_status}, Size: {len(docx_data)} bytes")
        
        # Test download mode
        handler = create_mock_handler(
            'generate-pdf',
            'POST',
            '/api/generate-pdf',
            json.dumps(test_document).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(test_document)))}
        )
        handler.do_POST()
        
        # Check response
        assert handler.response_status == 200
        docx_data = handler.response_body.getvalue()
        assert len(docx_data) > 0
        
        print(f"✓ PDF generation (download): Status {handler.response_status}, Size: {len(docx_data)} bytes")
        
        return True
        
    except Exception as e:
        print(f"✗ PDF generation endpoint test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_docx_conversion_endpoint():
    """Test the DOCX to PDF conversion endpoint."""
    print("\nTesting DOCX to PDF conversion endpoint...")
    
    try:
        # Create a test DOCX document
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document for Conversion', 0)
        doc.add_paragraph('This is a test document that will be converted to PDF.')
        doc.add_paragraph('It contains multiple paragraphs to test the conversion process.')
        
        # Save to memory
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_data = docx_buffer.getvalue()
        docx_buffer.close()
        
        # Encode as base64
        docx_base64 = base64.b64encode(docx_data).decode('utf-8')
        
        # Create conversion request
        conversion_request = {
            'docx_data': docx_base64
        }
        
        # Test preview mode
        handler = create_mock_handler(
            'convert-docx-pdf',
            'POST',
            '/api/convert-docx-pdf?preview=true',
            json.dumps(conversion_request).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(conversion_request)))}
        )
        handler.do_POST()
        
        # Check response
        assert handler.response_status == 200
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        
        assert response_json['success'] == True
        assert 'data' in response_json
        
        # Verify the base64 PDF data can be decoded
        pdf_data = base64.b64decode(response_json['data'])
        assert len(pdf_data) > 0
        assert pdf_data.startswith(b'%PDF-')  # PDF header
        
        print(f"✓ DOCX conversion (preview): Status {handler.response_status}, PDF Size: {len(pdf_data)} bytes")
        
        # Test download mode
        handler = create_mock_handler(
            'convert-docx-pdf',
            'POST',
            '/api/convert-docx-pdf',
            json.dumps(conversion_request).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(conversion_request)))}
        )
        handler.do_POST()
        
        # Check response
        assert handler.response_status == 200
        pdf_data = handler.response_body.getvalue()
        assert len(pdf_data) > 0
        assert pdf_data.startswith(b'%PDF-')  # PDF header
        
        print(f"✓ DOCX conversion (download): Status {handler.response_status}, PDF Size: {len(pdf_data)} bytes")
        
        return True
        
    except Exception as e:
        print(f"✗ DOCX conversion endpoint test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_error_handling():
    """Test error handling in the endpoints."""
    print("\nTesting error handling...")
    
    try:
        # Test invalid JSON
        handler = create_mock_handler(
            'generate-pdf',
            'POST',
            '/api/generate-pdf',
            b'invalid json',
            {'Content-Length': '12'}
        )
        handler.do_POST()
        
        assert handler.response_status == 400
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        assert response_json['success'] == False
        
        print("✓ Invalid JSON error handling works")
        
        # Test missing required fields
        invalid_document = {'title': 'Test'}  # Missing authors and sections
        
        handler = create_mock_handler(
            'generate-pdf',
            'POST',
            '/api/generate-pdf',
            json.dumps(invalid_document).encode('utf-8'),
            {'Content-Length': str(len(json.dumps(invalid_document)))}
        )
        handler.do_POST()
        
        assert handler.response_status == 400
        response_data = handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        assert response_json['success'] == False
        
        print("✓ Validation error handling works")
        
        return True
        
    except Exception as e:
        print(f"✗ Error handling test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run all endpoint tests."""
    print("Starting Python serverless endpoint tests...\n")
    
    tests = [
        test_health_endpoint,
        test_pdf_generation_endpoint,
        test_docx_conversion_endpoint,
        test_error_handling
    ]
    
    passed = 0
    total = len(tests)
    
    for test in tests:
        if test():
            passed += 1
    
    print(f"\n{'='*60}")
    print(f"Endpoint Test Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("✓ All endpoint tests passed! Python serverless functions are working correctly.")
        return 0
    else:
        print("✗ Some endpoint tests failed. Please check the errors above.")
        return 1

if __name__ == '__main__':
    sys.exit(main())
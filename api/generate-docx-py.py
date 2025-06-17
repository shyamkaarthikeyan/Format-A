import json
import sys
import os
from http.server import BaseHTTPRequestHandler
import io

# Copy your IEEE generator logic here
def generate_ieee_document(document_data):
    """
    Generate IEEE document using python-docx
    This should contain the logic from your ieee_generator_fixed.py
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
        
        # Create new document
        doc = Document()
        
        # Set up IEEE formatting
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        # Add title
        if document_data.get('title'):
            title = doc.add_heading(document_data['title'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add authors
        if document_data.get('authors'):
            authors_para = doc.add_paragraph()
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, author in enumerate(document_data['authors']):
                if i > 0:
                    authors_para.add_run(', ')
                authors_para.add_run(author.get('name', ''))
        
        # Add abstract
        if document_data.get('abstract'):
            abstract_para = doc.add_paragraph()
            abstract_run = abstract_para.add_run('Abstract—')
            abstract_run.italic = True
            abstract_para.add_run(document_data['abstract'])
        
        # Add keywords
        if document_data.get('keywords'):
            keywords_para = doc.add_paragraph()
            keywords_run = keywords_para.add_run('Index Terms—')
            keywords_run.italic = True
            keywords_para.add_run(document_data['keywords'])
        
        # Add sections
        if document_data.get('sections'):
            for i, section_data in enumerate(document_data['sections']):
                if section_data.get('title'):
                    heading = doc.add_heading(f"{i+1}. {section_data['title'].upper()}", level=1)
                
                if section_data.get('contentBlocks'):
                    for block in section_data['contentBlocks']:
                        if block.get('type') == 'text' and block.get('content'):
                            doc.add_paragraph(block['content'])
        
        # Add references
        if document_data.get('references'):
            doc.add_heading('REFERENCES', level=1)
            for i, ref in enumerate(document_data['references']):
                if ref.get('text'):
                    doc.add_paragraph(f"[{i+1}] {ref['text']}")
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        return doc_buffer.getvalue()
        
    except Exception as e:
        raise Exception(f"Document generation failed: {str(e)}")

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            document_data = json.loads(post_data.decode('utf-8'))
            
            # Generate document
            doc_bytes = generate_ieee_document(document_data)
            
            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename="ieee_paper.docx"')
            self.end_headers()
            self.wfile.write(doc_bytes)
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            error_response = json.dumps({'error': str(e)}).encode('utf-8')
            self.wfile.write(error_response)
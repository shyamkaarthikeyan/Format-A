from http.server import BaseHTTPRequestHandler
import json
import sys
import traceback
import time
from io import BytesIO
import base64

# Import utility functions for error handling and logging
try:
    from python_utils import (
        ServerlessLogger, ServerlessErrorHandler, PDFMetadataExtractor,
        DocumentValidator, TempFileManager, MemoryMonitor, TimeoutHandler,
        ProductionErrorHandler, VercelDiagnostics,
        create_success_response, send_json_response, send_binary_response
    )
except ImportError:
    # Fallback if utils not available
    class ServerlessLogger:
        @staticmethod
        def info(msg, **kwargs): print(f"INFO: {msg}", file=sys.stderr)
        @staticmethod
        def error(msg, **kwargs): print(f"ERROR: {msg}", file=sys.stderr)
        @staticmethod
        def warning(msg, **kwargs): print(f"WARNING: {msg}", file=sys.stderr)
    
    class ServerlessErrorHandler:
        @staticmethod
        def handle_exception(e, context="Operation"):
            return 500, {'success': False, 'error': str(e)}
    
    class PDFMetadataExtractor:
        @staticmethod
        def extract_basic_metadata(data):
            return {'size_bytes': len(data), 'generated_at': 'unknown'}
    
    class DocumentValidator:
        @staticmethod
        def validate_document_request(data): return True
    
    class MemoryMonitor:
        @staticmethod
        def check_memory_limit(): return True
        @staticmethod
        def get_memory_usage(): return {}
        @staticmethod
        def implement_graceful_degradation(info, op): return {'strategy': 'continue'}
    
    class TimeoutHandler:
        @staticmethod
        def check_timeout_risk(start, op): return {'at_risk': False}
        @staticmethod
        def create_timeout_response(elapsed, op): return {'error': 'timeout'}
    
    class ProductionErrorHandler:
        @staticmethod
        def create_production_error_response(e, ctx, req_id=None, debug=None):
            return {'success': False, 'error': str(e)}
    
    def create_success_response(data, message="Success", metadata=None):
        return {'success': True, 'data': data, 'message': message}
    
    def send_json_response(handler, status, data):
        handler.send_response(status)
        handler.send_header('Content-Type', 'application/json')
        handler.send_header('Access-Control-Allow-Origin', '*')
        handler.end_headers()
        handler.wfile.write(json.dumps(data).encode())

# Import all the existing IEEE generator functions
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from html.parser import HTMLParser
import unicodedata

def sanitize_text(text):
    """Sanitize text to remove invalid Unicode characters and surrogates."""
    if not text:
        return ""
    
    # Convert to string if not already
    text = str(text)
    
    # Remove surrogate characters and other problematic Unicode
    text = text.encode('utf-8', 'ignore').decode('utf-8')
    
    # Normalize Unicode characters
    text = unicodedata.normalize('NFKD', text)
    
    # Remove any remaining control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    return text

# IEEE formatting configuration - EXACT same as original
IEEE_CONFIG = {
    'font_name': 'Times New Roman',
    'font_size_title': Pt(24),
    'font_size_body': Pt(9.5),
    'font_size_caption': Pt(9),
    'margin_left': Inches(0.75),
    'margin_right': Inches(0.75),
    'margin_top': Inches(0.75),
    'margin_bottom': Inches(0.75),
    'column_count_body': 2,
    'column_spacing': Inches(0.25),
    'column_width': Inches(3.375),
    'column_indent': Inches(0.2),
    'line_spacing': Pt(10),  # Exact spacing for 9.5pt font
    'figure_sizes': {
        'Very Small': Inches(1.2),
        'Small': Inches(1.8),
        'Medium': Inches(2.5),
        'Large': Inches(3.2)
    },
    'max_figure_height': Inches(4.0),
}

def set_document_defaults(doc):
    """Set document-wide defaults to minimize unwanted spacing - EXACT same as original."""
    styles = doc.styles

    # Modify Normal style - EXACT same as original
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(12)
        normal.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        normal.paragraph_format.line_spacing_rule = 0  # Exact spacing
        normal.paragraph_format.widow_control = False
        normal.font.name = IEEE_CONFIG['font_name']
        normal.font.size = IEEE_CONFIG['font_size_body']
        # Add better spacing control
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Pt(0)

    # Modify Heading 1 style - LIKE ABSTRACT TITLE (regular weight, not bold)
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.base_style = styles['Normal']
        heading1.paragraph_format.space_before = Pt(0)
        heading1.paragraph_format.space_after = Pt(0)
        heading1.paragraph_format.line_spacing = Pt(10)
        heading1.paragraph_format.line_spacing_rule = 0
        heading1.paragraph_format.keep_with_next = False
        heading1.paragraph_format.page_break_before = False
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1.font.name = IEEE_CONFIG['font_name']
        heading1.font.size = IEEE_CONFIG['font_size_body']
        heading1.font.bold = False  # CHANGED: Like abstract title - NOT bold

    # Modify Heading 2 style for subsections - LIKE ABSTRACT TITLE (regular weight, not bold)
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.base_style = styles['Normal']
        heading2.paragraph_format.space_before = Pt(6)
        heading2.paragraph_format.space_after = Pt(0)
        heading2.paragraph_format.line_spacing = Pt(10)
        heading2.paragraph_format.line_spacing_rule = 0
        heading2.paragraph_format.keep_with_next = False
        heading2.paragraph_format.page_break_before = False
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading2.font.name = IEEE_CONFIG['font_name']
        heading2.font.size = IEEE_CONFIG['font_size_body']
        heading2.font.bold = False  # CHANGED: Like abstract title - NOT bold

def add_title(doc, title):
    """Add the paper title - EXACT same as original."""
    para = doc.add_paragraph()
    run = para.add_run(sanitize_text(title))
    run.bold = True
    run.font.name = IEEE_CONFIG['font_name']
    run.font.size = IEEE_CONFIG['font_size_title']
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)

def add_authors(doc, authors):
    """Add authors and their details in a parallel layout using a table - EXACT same as original."""
    if not authors:
        return
    
    num_authors = len(authors)
    table = doc.add_table(rows=1, cols=num_authors)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = True
    
    for idx, author in enumerate(authors):
        if not author.get('name'):
            continue
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        para = cell.add_paragraph()
        run = para.add_run(author['name'])
        run.bold = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        
        fields = [
            ('department', 'Department'),
            ('organization', 'Organization'),
            ('city', 'City'),
            ('state', 'State'),
            ('tamilnadu', 'Tamil Nadu')
        ]
        for field_key, field_name in fields:
            if author.get(field_key):
                para = cell.add_paragraph(sanitize_text(author[field_key]))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        for custom_field in author.get('custom_fields', []):
            if custom_field['value']:
                para = cell.add_paragraph(sanitize_text(custom_field['value']))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_abstract(doc, abstract):
    """Add the abstract section with bold title followed by content."""
    if abstract:
        # Add abstract with bold title and content in same paragraph
        para = doc.add_paragraph()
        
        # Bold "Abstract—" title (only bold, not italic)
        title_run = para.add_run("Abstract—")
        title_run.bold = True
        title_run.font.name = IEEE_CONFIG['font_name']
        title_run.font.size = IEEE_CONFIG['font_size_body']
        
        # Add abstract content immediately after (bold text)
        content_run = para.add_run(sanitize_text(abstract))
        content_run.bold = True
        content_run.font.name = IEEE_CONFIG['font_name']
        content_run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to abstract
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)

def add_keywords(doc, keywords):
    """Add the keywords section with bold title followed by content."""
    if keywords:
        # Add keywords with bold title and content in same paragraph
        para = doc.add_paragraph()
        
        # Bold "Keywords—" title (only bold, not italic)
        title_run = para.add_run("Keywords—")
        title_run.bold = True
        title_run.font.name = IEEE_CONFIG['font_name']
        title_run.font.size = IEEE_CONFIG['font_size_body']
        
        # Add keywords content immediately after (bold text)
        content_run = para.add_run(sanitize_text(keywords))
        content_run.bold = True
        content_run.font.name = IEEE_CONFIG['font_name']
        content_run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to keywords
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)
        
        # Minimal dummy paragraph to stabilize layout
        dummy_para = doc.add_paragraph("")
        dummy_para.paragraph_format.space_before = Pt(0)
        dummy_para.paragraph_format.space_after = Pt(0)
        dummy_para.paragraph_format.widow_control = False
        dummy_para.paragraph_format.keep_with_next = False
        dummy_para.paragraph_format.line_spacing = 0
        if dummy_para.runs:
            dummy_para.runs[0].font.size = Pt(1)

def add_justified_paragraph(doc, text, style_name='Normal', indent_left=None, indent_right=None, space_before=None, space_after=None):
    """Add a paragraph with optimized justification settings to prevent excessive word spacing - EXACT COPY from original."""
    para = doc.add_paragraph(sanitize_text(text))
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set paragraph formatting with exact spacing controls
    para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    para.paragraph_format.line_spacing_rule = 0  # Exact spacing
    para.paragraph_format.widow_control = False
    para.paragraph_format.keep_with_next = False
    para.paragraph_format.keep_together = False
    
    # Set spacing
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    
    # Set indentation
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    if indent_right is not None:
        para.paragraph_format.right_indent = indent_right
    
    # Font formatting with controlled spacing - RESTORED from original
    if para.runs:
        run = para.runs[0]
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Moderate character spacing controls (not aggressive) - ESSENTIAL for proper justification
        run_element = run._element
        rPr = run_element.get_or_add_rPr()
        
        # Set moderate character spacing to reduce word gaps without breaking words
        spacing_element = OxmlElement('w:spacing')
        spacing_element.set(qn('w:val'), '-5')  # Slight compression to reduce gaps
        rPr.append(spacing_element)
        
        # Prevent automatic text expansion but allow normal word flow
        run_element.set(qn('w:fitText'), '0')
    
    # Paragraph-level justification controls - MODERATE approach
    para_element = para._element
    pPr = para_element.get_or_add_pPr()
    
    # Use standard justification (not distribute) to keep words intact
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')  # Standard justify - keeps words together
    pPr.append(jc)
    
    # Control text alignment
    textAlignment = OxmlElement('w:textAlignment')
    textAlignment.set(qn('w:val'), 'baseline')
    pPr.append(textAlignment)
    
    # Moderate spacing control - prevent excessive gaps but allow normal flow
    adjust_right_ind = OxmlElement('w:adjustRightInd')
    adjust_right_ind.set(qn('w:val'), '0')
    pPr.append(adjust_right_ind)
    
    return para

class HTMLToWordParser(HTMLParser):
    """Parse HTML content and apply formatting to Word document."""
    
    def __init__(self, paragraph):
        super().__init__()
        self.paragraph = paragraph
        self.format_stack = []
        self.text_buffer = ""
    
    def handle_starttag(self, tag, attrs):
        # Flush any buffered text before starting new formatting
        self._flush_text()
        
        if tag.lower() in ['b', 'strong']:
            self.format_stack.append('bold')
        elif tag.lower() in ['i', 'em']:
            self.format_stack.append('italic')
        elif tag.lower() == 'u':
            self.format_stack.append('underline')
    
    def handle_endtag(self, tag):
        # Flush any buffered text before ending formatting
        self._flush_text()
        
        if tag.lower() in ['b', 'strong'] and 'bold' in self.format_stack:
            self.format_stack.remove('bold')
        elif tag.lower() in ['i', 'em'] and 'italic' in self.format_stack:
            self.format_stack.remove('italic')
        elif tag.lower() == 'u' and 'underline' in self.format_stack:
            self.format_stack.remove('underline')
    
    def handle_data(self, data):
        # Buffer the text data with sanitization
        self.text_buffer += sanitize_text(data)
    
    def _flush_text(self):
        """Create a run with accumulated text and current formatting."""
        if self.text_buffer:
            run = self.paragraph.add_run(self.text_buffer)
            run.font.name = IEEE_CONFIG['font_name']
            run.font.size = IEEE_CONFIG['font_size_body']
            
            # Apply current formatting
            if 'bold' in self.format_stack:
                run.bold = True
            if 'italic' in self.format_stack:
                run.italic = True
            if 'underline' in self.format_stack:
                run.underline = True
            
            self.text_buffer = ""
    
    def close(self):
        """Ensure any remaining text is flushed when parsing is complete."""
        self._flush_text()
        super().close()

def add_formatted_paragraph(doc, html_content, style_name='Normal', indent_left=None, indent_right=None, space_before=None, space_after=None):
    """Add a paragraph with HTML formatting support."""
    para = doc.add_paragraph(style=style_name)
    
    # Apply justification with advanced controls
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.widow_control = False
    para.paragraph_format.keep_with_next = False
    para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    para.paragraph_format.line_spacing_rule = 0  # Exact spacing
    
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    if indent_right is not None:
        para.paragraph_format.right_indent = indent_right
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    
    # Parse HTML and apply formatting
    if html_content and '<' in html_content and '>' in html_content:
        # Content contains HTML tags - use parser
        parser = HTMLToWordParser(para)
        parser.feed(html_content)
        parser.close()  # Important: flush any remaining text
    else:
        # Plain text content
        run = para.add_run(html_content or "")
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
    
    # Add advanced spacing controls to prevent word stretching
    para_element = para._element
    pPr = para_element.get_or_add_pPr()
    
    # Set justification method for better word spacing
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)
    
    return para

def add_section(doc, section_data, section_idx, is_first_section=False):
    """Add a section with content blocks, subsections, and figures - EXACT same as original."""
    if section_data.get('title'):
        para = doc.add_heading(f"{section_idx}. {sanitize_text(section_data['title']).upper()}", level=1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center section titles
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']  # Exactly one line before heading
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_together = False
        para.paragraph_format.widow_control = False

    # Process content blocks (text and images in order) - Support BOTH naming conventions
    content_blocks = section_data.get('contentBlocks', []) or section_data.get('content_blocks', [])
    
    for block_idx, block in enumerate(content_blocks):
        if block.get('type') == 'text' and block.get('content'):
            space_before = IEEE_CONFIG['line_spacing'] if is_first_section and block_idx == 0 else Pt(3)
            add_formatted_paragraph(
                doc, 
                block['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=space_before,
                space_after=Pt(12)
            )
            
            # Check if this text block also has an image attached (React frontend pattern)
            if block.get('data') and block.get('caption'):
                # Handle image attached to text block
                size = block.get('size', 'medium')
                # Map frontend size names to backend size names
                size_mapping = {
                    'very-small': 'Very Small',
                    'small': 'Small', 
                    'medium': 'Medium',
                    'large': 'Large'
                }
                mapped_size = size_mapping.get(size, 'Medium')
                width = IEEE_CONFIG['figure_sizes'].get(mapped_size, IEEE_CONFIG['figure_sizes']['Medium'])
                
                # Decode base64 image data
                try:
                    image_data = block['data']
                    
                    # Handle base64 data - remove prefix if present
                    if ',' in image_data:
                        image_data = image_data.split(',')[1]
                    
                    # Decode base64 image data
                    try:
                        image_bytes = base64.b64decode(image_data)
                    except Exception as e:
                        print(f"ERROR: Failed to decode image data in text block: {str(e)}", file=sys.stderr)
                        continue
                    
                    # Create image stream
                    image_stream = BytesIO(image_bytes)
                    
                    para = doc.add_paragraph()
                    run = para.add_run()
                    picture = run.add_picture(image_stream, width=width)
                    if picture.height > IEEE_CONFIG['max_figure_height']:
                        scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                        run.clear()
                        run.add_picture(image_stream, width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
                    
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    
                    # Generate figure number based on section and image position
                    img_count = sum(1 for b in content_blocks[:block_idx+1] if b.get('type') == 'image' or (b.get('type') == 'text' and b.get('data')))
                    caption = doc.add_paragraph(f"Fig. {section_idx}.{img_count}: {sanitize_text(block['caption'])}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_before = Pt(0)
                    caption.paragraph_format.space_after = Pt(12)
                    if caption.runs:
                        caption.runs[0].font.name = IEEE_CONFIG['font_name']
                        caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
                except Exception as e:
                    print(f"Error processing image in text block: {e}", file=sys.stderr)
                    
        elif block.get('type') == 'image' and block.get('data') and block.get('caption'):
            # Handle image blocks from React frontend
            size = block.get('size', 'medium')
            # Map frontend size names to backend size names
            size_mapping = {
                'very-small': 'Very Small',
                'small': 'Small', 
                'medium': 'Medium',
                'large': 'Large'
            }
            mapped_size = size_mapping.get(size, 'Medium')
            width = IEEE_CONFIG['figure_sizes'].get(mapped_size, IEEE_CONFIG['figure_sizes']['Medium'])
            
            # Decode base64 image data
            try:
                image_data = block['data']
                
                # Handle base64 data - remove prefix if present
                if ',' in image_data:
                    image_data = image_data.split(',')[1]
                
                # Decode base64 image data
                try:
                    image_bytes = base64.b64decode(image_data)
                except Exception as e:
                    print(f"ERROR: Failed to decode image data: {str(e)}", file=sys.stderr)
                    continue
                
                # Create image stream
                image_stream = BytesIO(image_bytes)
                
                para = doc.add_paragraph()
                run = para.add_run()
                picture = run.add_picture(image_stream, width=width)
                if picture.height > IEEE_CONFIG['max_figure_height']:
                    scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                    run.clear()
                    run.add_picture(image_stream, width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                # Generate figure number based on section and image position
                img_count = sum(1 for b in content_blocks[:block_idx+1] if b.get('type') == 'image')
                caption = doc.add_paragraph(f"Fig. {section_idx}.{img_count}: {sanitize_text(block['caption'])}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(0)
                caption.paragraph_format.space_after = Pt(12)
                if caption.runs:
                    caption.runs[0].font.name = IEEE_CONFIG['font_name']
                    caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
            except Exception as e:
                print(f"Error processing image: {e}", file=sys.stderr)

    # Legacy support for old content field - EXACT same as original
    if not content_blocks and section_data.get('content'):
        space_before = IEEE_CONFIG['line_spacing'] if is_first_section else Pt(3)
        add_justified_paragraph(
            doc, 
            section_data['content'],
            indent_left=IEEE_CONFIG['column_indent'],
            indent_right=IEEE_CONFIG['column_indent'],
            space_before=space_before,
            space_after=Pt(12)
        )

    # Add subsections with multi-level support
    def add_subsection_recursive(subsections, section_idx, parent_numbering=""):
        """Recursively add subsections with proper hierarchical numbering."""
        # Group subsections by level and parent
        level_1_subsections = [s for s in subsections if s.get('level', 1) == 1 and not s.get('parentId')]
        
        for sub_idx, subsection in enumerate(level_1_subsections, 1):
            if subsection.get('title'):
                subsection_number = f"{section_idx}.{sub_idx}"
                para = doc.add_heading(f"{subsection_number} {sanitize_text(subsection['title'])}", level=2)
                para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = False
                para.paragraph_format.widow_control = False

            if subsection.get('content'):
                add_justified_paragraph(
                    doc, 
                    sanitize_text(subsection['content']),
                    indent_left=IEEE_CONFIG['column_indent'],
                    indent_right=IEEE_CONFIG['column_indent'],
                    space_before=Pt(1),
                    space_after=Pt(12)
                )
            
            # Handle nested subsections (level 2 and beyond)
            add_nested_subsection(subsections, subsection['id'], f"{section_idx}.{sub_idx}", 2)
    
    def add_nested_subsection(all_subsections, parent_id, parent_number, level):
        """Add nested subsections recursively."""
        child_subsections = [s for s in all_subsections if s.get('parentId') == parent_id and s.get('level', 1) == level]
        
        for child_idx, child_sub in enumerate(child_subsections, 1):
            # Always define child_number, regardless of whether title exists
            child_number = f"{parent_number}.{child_idx}"
            
            if child_sub.get('title'):
                # Use different heading levels for deeper nesting, but cap at level 6
                heading_level = min(level + 1, 6)
                para = doc.add_heading(f"{child_number} {sanitize_text(child_sub['title'])}", level=heading_level)
                para.paragraph_format.page_break_before = False
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = False
                para.paragraph_format.widow_control = False

            if child_sub.get('content'):
                add_justified_paragraph(
                    doc, 
                    sanitize_text(child_sub['content']),
                    indent_left=IEEE_CONFIG['column_indent'] + Inches(0.1 * (level - 1)),  # Progressive indentation
                    indent_right=IEEE_CONFIG['column_indent'],
                    space_before=Pt(1),
                    space_after=Pt(12)
                )
            
            # Process content blocks if they exist
            if child_sub.get('contentBlocks'):
                for block in child_sub['contentBlocks']:
                    process_content_block(doc, block)
            
            # Recursively handle even deeper nesting
            if level < 5:  # Limit depth to prevent excessive nesting
                add_nested_subsection(all_subsections, child_sub['id'], child_number, level + 1)
    
    # Call the recursive function to add all subsections
    add_subsection_recursive(section_data.get('subsections', []), section_idx)

def add_references(doc, references):
    """Add references section."""
    if references:
        para = doc.add_heading("REFERENCES", level=1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']
        para.paragraph_format.space_after = Pt(0)
        
        for idx, ref in enumerate(references, 1):
            ref_para = doc.add_paragraph(f"[{idx}] {sanitize_text(ref)}")
            ref_para.paragraph_format.left_indent = IEEE_CONFIG['column_indent']
            ref_para.paragraph_format.space_before = Pt(3)
            ref_para.paragraph_format.space_after = Pt(6)
            if ref_para.runs:
                ref_para.runs[0].font.name = IEEE_CONFIG['font_name']
                ref_para.runs[0].font.size = IEEE_CONFIG['font_size_body']

def enable_auto_hyphenation(doc):
    """Enable automatic hyphenation for better text flow."""
    try:
        settings = doc.settings
        settings_element = settings._element
        
        # Add hyphenation settings
        compat = settings_element.xpath('./w:compat')
        if not compat:
            compat_element = OxmlElement('w:compat')
            settings_element.append(compat_element)
        else:
            compat_element = compat[0]
        
        # Enable auto hyphenation
        auto_hyphen = OxmlElement('w:autoHyphenation')
        auto_hyphen.set(qn('w:val'), '1')
        compat_element.append(auto_hyphen)
    except Exception as e:
        print(f"Warning: Could not enable auto-hyphenation: {e}", file=sys.stderr)

def set_compatibility_options(doc):
    """Set compatibility options for better rendering."""
    try:
        settings = doc.settings
        settings_element = settings._element
        
        # Add compatibility settings
        compat = settings_element.xpath('./w:compat')
        if not compat:
            compat_element = OxmlElement('w:compat')
            settings_element.append(compat_element)
        else:
            compat_element = compat[0]
        
        # Add various compatibility options for better rendering
        options = [
            ('balanceSingleByteDoubleByteWidth', '0'),
            ('doNotLeaveBackslashAlone', '0'),
            ('doNotExpandShiftReturn', '0'),
            ('adjustLineHeightInTable', '1'),
            ('useFELayout', '0')
        ]
        
        for option_name, option_value in options:
            option_element = OxmlElement(f'w:{option_name}')
            option_element.set(qn('w:val'), option_value)
            compat_element.append(option_element)
    except Exception as e:
        print(f"Warning: Could not set compatibility options: {e}", file=sys.stderr)

def generate_ieee_document(form_data):
    """Generate an IEEE-formatted Word document - EXACT same logic as original."""
    doc = Document()
    
    set_document_defaults(doc)
    
    section = doc.sections[0]
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    add_title(doc, form_data.get('title', ''))
    add_authors(doc, form_data.get('authors', []))

    # Add continuous section break for two-column layout
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.start_type = WD_SECTION.CONTINUOUS
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    # Set up the two-column layout FIRST before adding content
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    
    cols.set(qn('w:num'), str(IEEE_CONFIG['column_count_body']))
    cols.set(qn('w:sep'), '0')
    cols.set(qn('w:space'), str(int(IEEE_CONFIG['column_spacing'].pt)))
    cols.set(qn('w:equalWidth'), '1')
    
    # Add column definitions
    for i in range(IEEE_CONFIG['column_count_body']):
        col = OxmlElement('w:col')
        col.set(qn('w:w'), str(int(IEEE_CONFIG['column_width'].pt)))
        cols.append(col)
    
    # Prevent column balancing for stable layout
    no_balance = OxmlElement('w:noBalance')
    no_balance.set(qn('w:val'), '1')
    sectPr.append(no_balance)
    
    # Now add abstract and keywords in the properly configured 2-column layout
    add_abstract(doc, form_data.get('abstract', ''))
    add_keywords(doc, form_data.get('keywords', ''))
    
    for idx, section_data in enumerate(form_data.get('sections', []), 1):
        add_section(doc, section_data, idx, is_first_section=(idx == 1))
    
    add_references(doc, form_data.get('references', []))
    
    enable_auto_hyphenation(doc)
    set_compatibility_options(doc)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        request_id = f"pdf_gen_{id(self)}"
        start_time = time.time()
        
        ServerlessLogger.info("Starting PDF generation request", request_id=request_id)
        
        try:
            # Initial memory and timeout checks
            if not MemoryMonitor.check_memory_limit():
                error_response = ProductionErrorHandler.create_production_error_response(
                    MemoryError("Memory usage too high to start operation"),
                    "PDF generation initialization",
                    request_id=request_id
                )
                send_json_response(self, 507, error_response)
                return
            
            # Read request body with size validation
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 50 * 1024 * 1024:  # 50MB limit
                error_response = ProductionErrorHandler.create_production_error_response(
                    ValueError(f'Request size ({content_length / (1024*1024):.2f} MB) exceeds limit (50 MB)'),
                    "Request size validation",
                    request_id=request_id,
                    debug_info={'content_length': content_length, 'limit_mb': 50}
                )
                send_json_response(self, 413, error_response)
                return
            
            post_data = self.rfile.read(content_length)
            ServerlessLogger.info(f"Read request data: {len(post_data)} bytes", request_id=request_id)
            
            # Check timeout risk after reading data
            timeout_check = TimeoutHandler.check_timeout_risk(start_time, "pdf_generation")
            if timeout_check['at_risk'] and timeout_check['strategy'] == 'abort_gracefully':
                timeout_response = TimeoutHandler.create_timeout_response(
                    timeout_check['elapsed_seconds'], "pdf_generation"
                )
                send_json_response(self, 408, timeout_response)
                return
            
            # Parse JSON data
            try:
                form_data = json.loads(post_data.decode('utf-8'))
                ServerlessLogger.info("Successfully parsed JSON request", request_id=request_id)
            except json.JSONDecodeError as e:
                error_response = ProductionErrorHandler.create_production_error_response(
                    e, "JSON parsing", request_id=request_id,
                    debug_info={'data_length': len(post_data), 'data_preview': post_data[:100].decode('utf-8', errors='ignore')}
                )
                send_json_response(self, 400, error_response)
                return
            
            # Validate document request
            try:
                DocumentValidator.validate_document_request(form_data)
                ServerlessLogger.info("Document request validation passed", request_id=request_id)
            except ValueError as e:
                error_response = ProductionErrorHandler.create_production_error_response(
                    e, "Document validation", request_id=request_id,
                    debug_info={'form_data_keys': list(form_data.keys()) if isinstance(form_data, dict) else 'not_dict'}
                )
                send_json_response(self, 400, error_response)
                return
            
            # Check memory before generation
            memory_info = MemoryMonitor.get_memory_usage()
            degradation_strategy = MemoryMonitor.implement_graceful_degradation(memory_info, "pdf_generation")
            
            if degradation_strategy['strategy'] == 'abort':
                error_response = ProductionErrorHandler.create_production_error_response(
                    MemoryError(degradation_strategy['message']),
                    "Memory limit check",
                    request_id=request_id,
                    debug_info={'memory_info': memory_info, 'degradation_strategy': degradation_strategy}
                )
                send_json_response(self, 507, error_response)
                return
            elif degradation_strategy['strategy'] in ['reduce_quality', 'optimize']:
                ServerlessLogger.warning(
                    f"Applying degradation strategy: {degradation_strategy['strategy']}",
                    request_id=request_id,
                    strategy=degradation_strategy
                )
            
            # Generate IEEE document with timeout monitoring
            try:
                ServerlessLogger.info("Starting document generation", request_id=request_id)
                
                # Check timeout before starting generation
                timeout_check = TimeoutHandler.check_timeout_risk(start_time, "pdf_generation")
                if timeout_check['at_risk'] and timeout_check['strategy'] == 'expedite_operation':
                    ServerlessLogger.warning("Expediting operation due to timeout risk", request_id=request_id)
                
                doc_data = generate_ieee_document(form_data)
                
                # Check timeout after generation
                timeout_check = TimeoutHandler.check_timeout_risk(start_time, "pdf_generation")
                if timeout_check['at_risk']:
                    ServerlessLogger.warning("Operation completed but close to timeout", 
                                           request_id=request_id, timeout_info=timeout_check)
                
                ServerlessLogger.info(f"Document generated successfully: {len(doc_data)} bytes", request_id=request_id)
                
                # Extract metadata with performance info
                metadata = PDFMetadataExtractor.extract_basic_metadata(doc_data)
                metadata['request_id'] = request_id
                metadata['generation_time_seconds'] = round(time.time() - start_time, 2)
                metadata['memory_usage'] = MemoryMonitor.get_memory_usage()
                
                # Log final memory usage
                final_memory = MemoryMonitor.get_memory_usage()
                if 'rss_mb' in final_memory:
                    ServerlessLogger.info(f"Memory usage after generation: {final_memory['rss_mb']} MB", request_id=request_id)
                
                # Determine response format based on query parameters
                query_params = self.path.split('?', 1)
                is_preview = False
                if len(query_params) > 1:
                    import urllib.parse
                    params = urllib.parse.parse_qs(query_params[1])
                    is_preview = params.get('preview', ['false'])[0].lower() == 'true'
                
                if is_preview:
                    # Return base64 encoded data for preview
                    doc_base64 = base64.b64encode(doc_data).decode('utf-8')
                    response = create_success_response(
                        doc_base64,
                        'IEEE document generated successfully for preview',
                        metadata
                    )
                    response['format'] = 'docx'
                    
                    ServerlessLogger.info("Sending preview response", request_id=request_id)
                    send_json_response(self, 200, response)
                else:
                    # Return binary data for download
                    ServerlessLogger.info("Sending binary download response", request_id=request_id)
                    send_binary_response(self, doc_data, 
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'ieee_document.docx')
                    
            except Exception as e:
                # Enhanced error handling with production debugging
                elapsed_time = time.time() - start_time
                debug_info = {
                    'elapsed_seconds': elapsed_time,
                    'memory_info': MemoryMonitor.get_memory_usage(),
                    'timeout_info': TimeoutHandler.check_timeout_risk(start_time, "pdf_generation"),
                    'form_data_summary': {
                        'title': form_data.get('title', 'missing')[:50] if form_data.get('title') else 'missing',
                        'authors_count': len(form_data.get('authors', [])),
                        'sections_count': len(form_data.get('sections', []))
                    }
                }
                
                error_response = ProductionErrorHandler.create_production_error_response(
                    e, "Document generation", request_id=request_id, debug_info=debug_info
                )
                send_json_response(self, 500, error_response)
                
        except Exception as e:
            # Top-level error handling with full diagnostics
            elapsed_time = time.time() - start_time
            debug_info = {
                'elapsed_seconds': elapsed_time,
                'memory_info': MemoryMonitor.get_memory_usage(),
                'diagnostics': VercelDiagnostics.get_function_diagnostics()
            }
            
            error_response = ProductionErrorHandler.create_production_error_response(
                e, "PDF generation request processing", request_id=request_id, debug_info=debug_info
            )
            send_json_response(self, 500, error_response)
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.end_headers()
    
    def send_error_response(self, status_code, error_code, message, details=None):
        """Send a structured error response."""
        response = {
            'success': False,
            'error': {
                'code': error_code,
                'message': message,
                'timestamp': sys.version,
                'environment': 'vercel_serverless'
            }
        }
        
        if details:
            response['error']['details'] = details
            # Log detailed error to stderr for Vercel logs
            print(f"ERROR: {message}", file=sys.stderr)
            print(f"DETAILS: {details}", file=sys.stderr)
        
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(response, indent=2).encode())
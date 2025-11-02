from http.server import BaseHTTPRequestHandler
import json
import sys
import traceback
import os
import tempfile
import time
from io import BytesIO
import base64

# Import enhanced utility functions for error handling and logging
try:
    from python_utils import (
        ServerlessLogger, ServerlessErrorHandler, PDFMetadataExtractor,
        DocumentValidator, TempFileManager, MemoryMonitor, TimeoutHandler,
        ProductionErrorHandler, VercelDiagnostics,
        create_success_response, send_json_response, send_binary_response
    )
except ImportError:
    # Fallback if utils not available
    class ServerlessLogger:
        @staticmethod
        def info(msg, **kwargs): print(f"INFO: {msg}", file=sys.stderr)
        @staticmethod
        def error(msg, **kwargs): print(f"ERROR: {msg}", file=sys.stderr)
        @staticmethod
        def warning(msg, **kwargs): print(f"WARNING: {msg}", file=sys.stderr)
    
    class MemoryMonitor:
        @staticmethod
        def get_memory_usage(): return {}
    
    class TimeoutHandler:
        @staticmethod
        def check_timeout_risk(start, op): return {'at_risk': False}
    
    class VercelDiagnostics:
        @staticmethod
        def get_function_diagnostics(): return {}
    
    class ProductionErrorHandler:
        @staticmethod
        def create_production_error_response(e, ctx, req_id=None, debug=None):
            return {'success': False, 'error': str(e)}

class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        """Health check endpoint for Python serverless functions with enhanced diagnostics."""
        request_id = f"health_check_{id(self)}"
        start_time = time.time()
        
        ServerlessLogger.info("Starting health check request", request_id=request_id)
        
        try:
            # Get comprehensive diagnostics including memory and timeout info
            health_data = self.check_python_health()
            
            # Add serverless-specific diagnostics
            memory_info = MemoryMonitor.get_memory_usage()
            timeout_info = TimeoutHandler.check_timeout_risk(start_time, "health_check")
            vercel_diagnostics = VercelDiagnostics.get_function_diagnostics()
            
            # Determine overall health status
            all_healthy = all(dep['status'] == 'available' for dep in health_data['dependencies'])
            
            # Check for memory or timeout issues
            memory_healthy = True
            if 'rss_mb' in memory_info and memory_info['rss_mb'] > 400:  # 80% of 512MB limit
                memory_healthy = False
                ServerlessLogger.warning(f"High memory usage detected: {memory_info['rss_mb']} MB", request_id=request_id)
            
            timeout_healthy = not timeout_info.get('at_risk', False)
            
            overall_healthy = all_healthy and memory_healthy and timeout_healthy
            
            response = {
                'success': True,
                'status': 'healthy' if overall_healthy else 'degraded',
                'timestamp': self.get_timestamp(),
                'request_id': request_id,
                'execution_time_seconds': round(time.time() - start_time, 3),
                'environment': self.get_environment_info(),
                'dependencies': health_data['dependencies'],
                'capabilities': health_data['capabilities'],
                'system': health_data['system'],
                'serverless_diagnostics': {
                    'memory_info': memory_info,
                    'timeout_info': timeout_info,
                    'vercel_diagnostics': vercel_diagnostics,
                    'memory_healthy': memory_healthy,
                    'timeout_healthy': timeout_healthy
                }
            }
            
            # Add warnings if any issues detected
            warnings = []
            if not memory_healthy:
                warnings.append("High memory usage detected - may affect performance")
            if not timeout_healthy:
                warnings.append("Function execution time approaching limits")
            if not all_healthy:
                missing_deps = [dep['name'] for dep in health_data['dependencies'] if dep['status'] != 'available']
                warnings.append(f"Missing dependencies: {', '.join(missing_deps)}")
            
            if warnings:
                response['warnings'] = warnings
            
            status_code = 200 if overall_healthy else 206  # 206 for partial content/degraded
            
            ServerlessLogger.info(f"Health check completed: {response['status']}", 
                                request_id=request_id, 
                                execution_time=response['execution_time_seconds'])
            
            self.send_response(status_code)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(response, indent=2).encode())
            
        except Exception as e:
            # Enhanced error handling for health check failures
            elapsed_time = time.time() - start_time
            debug_info = {
                'elapsed_seconds': elapsed_time,
                'memory_info': MemoryMonitor.get_memory_usage(),
                'diagnostics': VercelDiagnostics.get_function_diagnostics()
            }
            
            error_response = ProductionErrorHandler.create_production_error_response(
                e, "Health check execution", request_id=request_id, debug_info=debug_info
            )
            
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(error_response, indent=2).encode())
    
    def do_POST(self):
        """Detailed health check with optional dependency testing."""
        try:
            # Read request body for test parameters
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 0:
                post_data = self.rfile.read(content_length)
                try:
                    request_data = json.loads(post_data.decode('utf-8'))
                except json.JSONDecodeError:
                    request_data = {}
            else:
                request_data = {}
            
            # Perform detailed health check
            health_data = self.check_python_health()
            
            # Run additional tests if requested
            test_pdf_generation = request_data.get('test_pdf_generation', False)
            test_docx_conversion = request_data.get('test_docx_conversion', False)
            
            if test_pdf_generation:
                health_data['tests'] = health_data.get('tests', {})
                health_data['tests']['pdf_generation'] = self.test_pdf_generation()
            
            if test_docx_conversion:
                health_data['tests'] = health_data.get('tests', {})
                health_data['tests']['docx_conversion'] = self.test_docx_conversion()
            
            # Determine overall health status
            all_healthy = all(dep['status'] == 'available' for dep in health_data['dependencies'])
            if 'tests' in health_data:
                all_tests_passed = all(test.get('status') == 'passed' for test in health_data['tests'].values())
                all_healthy = all_healthy and all_tests_passed
            
            response = {
                'success': True,
                'status': 'healthy' if all_healthy else 'degraded',
                'timestamp': self.get_timestamp(),
                'environment': self.get_environment_info(),
                'dependencies': health_data['dependencies'],
                'capabilities': health_data['capabilities'],
                'system': health_data['system']
            }
            
            if 'tests' in health_data:
                response['tests'] = health_data['tests']
            
            status_code = 200 if all_healthy else 206
            
            self.send_response(status_code)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(response, indent=2).encode())
            
        except Exception as e:
            self.send_error_response(500, 'HEALTH_CHECK_FAILED', f'Detailed health check failed: {str(e)}', traceback.format_exc())
    
    def check_python_health(self):
        """Check the health of Python dependencies and environment."""
        dependencies = []
        capabilities = []
        
        # Check reportlab
        reportlab_status = self.check_dependency('reportlab', 'reportlab.pdfgen.canvas')
        dependencies.append(reportlab_status)
        if reportlab_status['status'] == 'available':
            capabilities.append('pdf_generation')
        
        # Check python-docx
        docx_status = self.check_dependency('python-docx', 'docx')
        dependencies.append(docx_status)
        if docx_status['status'] == 'available':
            capabilities.append('docx_processing')
        
        # Check Pillow
        pillow_status = self.check_dependency('Pillow', 'PIL.Image')
        dependencies.append(pillow_status)
        if pillow_status['status'] == 'available':
            capabilities.append('image_processing')
        
        # Check docx2pdf
        docx2pdf_status = self.check_dependency('docx2pdf', 'docx2pdf')
        dependencies.append(docx2pdf_status)
        if docx2pdf_status['status'] == 'available':
            capabilities.append('docx_to_pdf_conversion')
        
        # System information
        system_info = {
            'python_version': sys.version,
            'platform': sys.platform,
            'temp_directory': self.get_temp_directory_info(),
            'memory_info': self.get_memory_info(),
            'environment_type': 'vercel_serverless'
        }
        
        return {
            'dependencies': dependencies,
            'capabilities': capabilities,
            'system': system_info
        }
    
    def check_dependency(self, package_name, import_path):
        """Check if a specific dependency is available."""
        try:
            # Try to import the module
            __import__(import_path)
            
            # Try to get version if possible
            version = 'unknown'
            try:
                if package_name == 'reportlab':
                    import reportlab
                    version = getattr(reportlab, '__version__', 'unknown')
                elif package_name == 'python-docx':
                    import docx
                    version = getattr(docx, '__version__', 'unknown')
                elif package_name == 'Pillow':
                    import PIL
                    version = getattr(PIL, '__version__', 'unknown')
                elif package_name == 'docx2pdf':
                    # docx2pdf doesn't have __version__, try to get it from package info
                    try:
                        import pkg_resources
                        version = pkg_resources.get_distribution('docx2pdf').version
                    except:
                        version = 'available'
            except:
                pass
            
            return {
                'name': package_name,
                'status': 'available',
                'version': version,
                'import_path': import_path
            }
            
        except ImportError as e:
            return {
                'name': package_name,
                'status': 'missing',
                'error': str(e),
                'import_path': import_path
            }
        except Exception as e:
            return {
                'name': package_name,
                'status': 'error',
                'error': str(e),
                'import_path': import_path
            }
    
    def test_pdf_generation(self):
        """Test basic PDF generation capability."""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Create a test PDF in memory
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, "Health Check Test PDF")
            c.drawString(100, 730, "This PDF was generated successfully by the serverless function.")
            c.save()
            
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return {
                'status': 'passed',
                'message': 'PDF generation test successful',
                'pdf_size': len(pdf_data),
                'test_type': 'basic_pdf_generation'
            }
            
        except Exception as e:
            return {
                'status': 'failed',
                'message': f'PDF generation test failed: {str(e)}',
                'error': str(e),
                'test_type': 'basic_pdf_generation'
            }
    
    def test_docx_conversion(self):
        """Test DOCX to PDF conversion capability."""
        try:
            # First create a simple DOCX file
            from docx import Document
            
            doc = Document()
            doc.add_heading('Health Check Test Document', 0)
            doc.add_paragraph('This is a test document created for health check purposes.')
            
            # Save DOCX to memory
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            # Try to convert to PDF using docx2pdf
            try:
                from docx2pdf import convert
                
                # Create temporary files
                temp_dir = '/tmp'
                if not os.path.exists(temp_dir):
                    temp_dir = tempfile.gettempdir()
                
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, dir=temp_dir) as temp_docx:
                    temp_docx.write(docx_data)
                    temp_docx_path = temp_docx.name
                
                temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                
                try:
                    convert(temp_docx_path, temp_pdf_path)
                    
                    if os.path.exists(temp_pdf_path):
                        pdf_size = os.path.getsize(temp_pdf_path)
                        
                        # Clean up
                        os.unlink(temp_docx_path)
                        os.unlink(temp_pdf_path)
                        
                        return {
                            'status': 'passed',
                            'message': 'DOCX to PDF conversion test successful',
                            'docx_size': len(docx_data),
                            'pdf_size': pdf_size,
                            'test_type': 'docx_to_pdf_conversion'
                        }
                    else:
                        return {
                            'status': 'failed',
                            'message': 'DOCX to PDF conversion failed - no output file created',
                            'test_type': 'docx_to_pdf_conversion'
                        }
                        
                finally:
                    # Ensure cleanup
                    try:
                        if os.path.exists(temp_docx_path):
                            os.unlink(temp_docx_path)
                        if os.path.exists(temp_pdf_path):
                            os.unlink(temp_pdf_path)
                    except:
                        pass
                        
            except ImportError:
                return {
                    'status': 'skipped',
                    'message': 'docx2pdf not available - conversion test skipped',
                    'test_type': 'docx_to_pdf_conversion'
                }
                
        except Exception as e:
            return {
                'status': 'failed',
                'message': f'DOCX conversion test failed: {str(e)}',
                'error': str(e),
                'test_type': 'docx_to_pdf_conversion'
            }
    
    def get_temp_directory_info(self):
        """Get information about the temporary directory."""
        try:
            temp_dirs = ['/tmp', tempfile.gettempdir()]
            temp_info = {}
            
            for temp_dir in temp_dirs:
                if os.path.exists(temp_dir):
                    temp_info[temp_dir] = {
                        'exists': True,
                        'writable': os.access(temp_dir, os.W_OK),
                        'readable': os.access(temp_dir, os.R_OK)
                    }
                    
                    # Try to create a test file
                    try:
                        test_file = os.path.join(temp_dir, 'health_check_test.tmp')
                        with open(test_file, 'w') as f:
                            f.write('test')
                        os.unlink(test_file)
                        temp_info[temp_dir]['test_write'] = True
                    except:
                        temp_info[temp_dir]['test_write'] = False
                else:
                    temp_info[temp_dir] = {'exists': False}
            
            return temp_info
            
        except Exception as e:
            return {'error': str(e)}
    
    def get_memory_info(self):
        """Get basic memory information if available."""
        try:
            import psutil
            memory = psutil.virtual_memory()
            return {
                'total': memory.total,
                'available': memory.available,
                'percent': memory.percent,
                'used': memory.used
            }
        except ImportError:
            return {'error': 'psutil not available'}
        except Exception as e:
            return {'error': str(e)}
    
    def get_environment_info(self):
        """Get environment information."""
        return {
            'vercel_region': os.environ.get('VERCEL_REGION', 'unknown'),
            'vercel_env': os.environ.get('VERCEL_ENV', 'unknown'),
            'node_env': os.environ.get('NODE_ENV', 'unknown'),
            'python_path': sys.executable,
            'working_directory': os.getcwd()
        }
    
    def get_timestamp(self):
        """Get current timestamp."""
        try:
            from datetime import datetime
            return datetime.utcnow().isoformat() + 'Z'
        except:
            import time
            return str(time.time())
    
    def do_OPTIONS(self):
        """Handle CORS preflight requests."""
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.end_headers()
    
    def send_error_response(self, status_code, error_code, message, details=None):
        """Send a structured error response."""
        response = {
            'success': False,
            'status': 'unhealthy',
            'error': {
                'code': error_code,
                'message': message,
                'timestamp': self.get_timestamp(),
                'environment': 'vercel_serverless'
            }
        }
        
        if details:
            response['error']['details'] = details
            # Log detailed error to stderr for Vercel logs
            print(f"HEALTH CHECK ERROR: {message}", file=sys.stderr)
            print(f"DETAILS: {details}", file=sys.stderr)
        
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(response, indent=2).encode())
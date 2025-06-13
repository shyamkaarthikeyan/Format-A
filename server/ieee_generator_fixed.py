"""
IEEE Document Generator - EXACT copy from test.py 
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# IEEE formatting configuration - EXACT same as test.py
IEEE_CONFIG = {
    'font_name': 'Times New Roman',
    'font_size_title': Pt(24),
    'font_size_body': Pt(9.5),
    'font_size_caption': Pt(9),
    'margin_left': Inches(0.75),
    'margin_right': Inches(0.75),
    'margin_top': Inches(0.75),
    'margin_bottom': Inches(0.75),
    'column_count_body': 2,
    'column_spacing': Inches(0.25),
    'column_width': Inches(3.375),
    'column_indent': Inches(0.2),
    'line_spacing': Pt(10),  # Exact spacing for 9.5pt font
    'figure_sizes': {
        'Very Small': Inches(1.2),
        'Small': Inches(1.8),
        'Medium': Inches(2.5),
        'Large': Inches(3.2)
    },
    'max_figure_height': Inches(4.0),
}

def set_document_defaults(doc):
    """Set document-wide defaults to minimize unwanted spacing - EXACT same as test.py."""
    styles = doc.styles

    # Modify Normal style - EXACT same as test.py
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(12)
        normal.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        normal.paragraph_format.line_spacing_rule = 0  # Exact spacing
        normal.paragraph_format.widow_control = False
        normal.font.name = IEEE_CONFIG['font_name']
        normal.font.size = IEEE_CONFIG['font_size_body']
        # Add better spacing control
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Pt(0)

    # Modify Heading 1 style - LIKE ABSTRACT TITLE (regular weight, not bold)
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.base_style = styles['Normal']
        heading1.paragraph_format.space_before = Pt(0)
        heading1.paragraph_format.space_after = Pt(0)
        heading1.paragraph_format.line_spacing = Pt(10)
        heading1.paragraph_format.line_spacing_rule = 0
        heading1.paragraph_format.keep_with_next = False
        heading1.paragraph_format.page_break_before = False
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1.font.name = IEEE_CONFIG['font_name']
        heading1.font.size = IEEE_CONFIG['font_size_body']
        heading1.font.bold = False  # CHANGED: Like abstract title - NOT bold

    # Modify Heading 2 style for subsections - LIKE ABSTRACT TITLE (regular weight, not bold)
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.base_style = styles['Normal']
        heading2.paragraph_format.space_before = Pt(6)
        heading2.paragraph_format.space_after = Pt(0)
        heading2.paragraph_format.line_spacing = Pt(10)
        heading2.paragraph_format.line_spacing_rule = 0
        heading2.paragraph_format.keep_with_next = False
        heading2.paragraph_format.page_break_before = False
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading2.font.name = IEEE_CONFIG['font_name']
        heading2.font.size = IEEE_CONFIG['font_size_body']
        heading2.font.bold = False  # CHANGED: Like abstract title - NOT bold

def add_title(doc, title):
    """Add the paper title - EXACT same as test.py."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = IEEE_CONFIG['font_name']
    run.font.size = IEEE_CONFIG['font_size_title']
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)

def add_authors(doc, authors):
    """Add authors and their details in a parallel layout using a table - EXACT same as test.py."""
    if not authors:
        return
    
    num_authors = len(authors)
    table = doc.add_table(rows=1, cols=num_authors)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = True
    
    for idx, author in enumerate(authors):
        if not author.get('name'):
            continue
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        para = cell.add_paragraph()
        run = para.add_run(author['name'])
        run.bold = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        
        fields = [
            ('department', 'Department'),
            ('organization', 'Organization'),
            ('city', 'City'),
            ('state', 'State'),
            ('tamilnadu', 'Tamil Nadu')
        ]
        for field_key, field_name in fields:
            if author.get(field_key):
                para = cell.add_paragraph(author[field_key])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        for custom_field in author.get('custom_fields', []):
            if custom_field['value']:
                para = cell.add_paragraph(custom_field['value'])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_abstract(doc, abstract):
    """Add the abstract section with italicized 'Abstract—' - EXACT same as test.py."""
    if abstract:
        para = doc.add_paragraph()
        run = para.add_run("Abstract—")
        run.italic = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        run = para.add_run(abstract)
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to abstract - EXACT same as test.py
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching - EXACT same as test.py
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)

def add_keywords(doc, keywords):
    """Add the keywords section - EXACT same as test.py."""
    if keywords:
        para = doc.add_paragraph(f"Keywords: {keywords}")
        
        # Apply advanced justification controls to keywords - EXACT same as test.py
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        if para.runs:
            para.runs[0].font.name = IEEE_CONFIG['font_name']
            para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        # Add advanced spacing controls to prevent word stretching - EXACT same as test.py
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)
        
        # Minimal dummy paragraph to stabilize layout - EXACT same as test.py
        dummy_para = doc.add_paragraph("")
        dummy_para.paragraph_format.space_before = Pt(0)
        dummy_para.paragraph_format.space_after = Pt(0)
        dummy_para.paragraph_format.widow_control = False
        dummy_para.paragraph_format.keep_with_next = False
        dummy_para.paragraph_format.line_spacing = 0
        if dummy_para.runs:
            dummy_para.runs[0].font.size = Pt(1)

def add_justified_paragraph(doc, text, style_name='Normal', indent_left=None, indent_right=None, space_before=None, space_after=None):
    """Add a paragraph with optimized justification settings to prevent excessive word spacing - EXACT COPY from test.py."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set paragraph formatting with exact spacing controls
    para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    para.paragraph_format.line_spacing_rule = 0  # Exact spacing
    para.paragraph_format.widow_control = False
    para.paragraph_format.keep_with_next = False
    para.paragraph_format.keep_together = False
    
    # Set spacing
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    
    # Set indentation
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    if indent_right is not None:
        para.paragraph_format.right_indent = indent_right
    
    # Font formatting with controlled spacing - RESTORED from test.py
    if para.runs:
        run = para.runs[0]
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Moderate character spacing controls (not aggressive) - ESSENTIAL for proper justification
        run_element = run._element
        rPr = run_element.get_or_add_rPr()
        
        # Set moderate character spacing to reduce word gaps without breaking words
        spacing_element = OxmlElement('w:spacing')
        spacing_element.set(qn('w:val'), '-5')  # Slight compression to reduce gaps
        rPr.append(spacing_element)
        
        # Prevent automatic text expansion but allow normal word flow
        run_element.set(qn('w:fitText'), '0')
    
    # Paragraph-level justification controls - MODERATE approach
    para_element = para._element
    pPr = para_element.get_or_add_pPr()
    
    # Use standard justification (not distribute) to keep words intact
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')  # Standard justify - keeps words together
    pPr.append(jc)
    
    # Control text alignment
    textAlignment = OxmlElement('w:textAlignment')
    textAlignment.set(qn('w:val'), 'baseline')
    pPr.append(textAlignment)
    
    # Moderate spacing control - prevent excessive gaps but allow normal flow
    adjust_right_ind = OxmlElement('w:adjustRightInd')
    adjust_right_ind.set(qn('w:val'), '0')
    pPr.append(adjust_right_ind)
    
    return para

def add_section(doc, section_data, section_idx, is_first_section=False):
    """Add a section with content blocks, subsections, and figures - EXACT same as test.py."""
    if section_data.get('title'):
        para = doc.add_heading(f"{section_idx}. {section_data['title'].upper()}", level=1)
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']  # Exactly one line before heading
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_together = False
        para.paragraph_format.widow_control = False

    # Process content blocks (text and images in order) - Support BOTH naming conventions
    content_blocks = section_data.get('contentBlocks', []) or section_data.get('content_blocks', [])
    
    for block_idx, block in enumerate(content_blocks):
        if block.get('type') == 'text' and block.get('content'):
            space_before = IEEE_CONFIG['line_spacing'] if is_first_section and block_idx == 0 else Pt(3)
            add_justified_paragraph(
                doc, 
                block['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=space_before,
                space_after=Pt(12)
            )
            
            # Check if this text block also has an image attached (React frontend pattern)
            if block.get('data') and block.get('caption'):
                # Handle image attached to text block
                import base64
                size = block.get('size', 'medium')
                # Map frontend size names to backend size names
                size_mapping = {
                    'very-small': 'Very Small',
                    'small': 'Small', 
                    'medium': 'Medium',
                    'large': 'Large'
                }
                mapped_size = size_mapping.get(size, 'Medium')
                width = IEEE_CONFIG['figure_sizes'].get(mapped_size, IEEE_CONFIG['figure_sizes']['Medium'])
                
                # Decode base64 image data
                try:
                    image_data = block['data']
                    
                    # Handle base64 data - remove prefix if present
                    if ',' in image_data:
                        image_data = image_data.split(',')[1]
                    
                    # Decode base64 image data
                    try:
                        image_bytes = base64.b64decode(image_data)
                    except Exception as e:
                        print(f"ERROR: Failed to decode image data in text block: {str(e)}", file=sys.stderr)
                        continue
                    
                    # Create image stream
                    image_stream = BytesIO(image_bytes)
                    
                    para = doc.add_paragraph()
                    run = para.add_run()
                    picture = run.add_picture(image_stream, width=width)
                    if picture.height > IEEE_CONFIG['max_figure_height']:
                        scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                        run.clear()
                        run.add_picture(image_stream, width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
                    
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    
                    # Generate figure number based on section and image position
                    img_count = sum(1 for b in content_blocks[:block_idx+1] if b.get('type') == 'image' or (b.get('type') == 'text' and b.get('data')))
                    caption = doc.add_paragraph(f"Fig. {section_idx}.{img_count}: {block['caption']}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_before = Pt(0)
                    caption.paragraph_format.space_after = Pt(12)
                    if caption.runs:
                        caption.runs[0].font.name = IEEE_CONFIG['font_name']
                        caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
                except Exception as e:
                    print(f"Error processing image in text block: {e}", file=sys.stderr)
                    
        elif block.get('type') == 'image' and block.get('data') and block.get('caption'):
            # Handle image blocks from React frontend
            import base64
            size = block.get('size', 'medium')
            # Map frontend size names to backend size names
            size_mapping = {
                'very-small': 'Very Small',
                'small': 'Small', 
                'medium': 'Medium',
                'large': 'Large'
            }
            mapped_size = size_mapping.get(size, 'Medium')
            width = IEEE_CONFIG['figure_sizes'].get(mapped_size, IEEE_CONFIG['figure_sizes']['Medium'])
            
            # Decode base64 image data
            try:
                image_data = block['data']
                
                # Handle base64 data - remove prefix if present
                if ',' in image_data:
                    image_data = image_data.split(',')[1]
                
                # Decode base64 image data
                try:
                    image_bytes = base64.b64decode(image_data)
                except Exception as e:
                    print(f"ERROR: Failed to decode image data: {str(e)}", file=sys.stderr)
                    continue
                
                # Create image stream
                image_stream = BytesIO(image_bytes)
                
                para = doc.add_paragraph()
                run = para.add_run()
                picture = run.add_picture(image_stream, width=width)
                if picture.height > IEEE_CONFIG['max_figure_height']:
                    scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                    run.clear()
                    run.add_picture(image_stream, width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                # Generate figure number based on section and image position
                img_count = sum(1 for b in content_blocks[:block_idx+1] if b.get('type') == 'image')
                caption = doc.add_paragraph(f"Fig. {section_idx}.{img_count}: {block['caption']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(0)
                caption.paragraph_format.space_after = Pt(12)
                if caption.runs:
                    caption.runs[0].font.name = IEEE_CONFIG['font_name']
                    caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
            except Exception as e:
                print(f"Error processing image: {e}", file=sys.stderr)

    # Legacy support for old content field - EXACT same as test.py
    if not content_blocks and section_data.get('content'):
        space_before = IEEE_CONFIG['line_spacing'] if is_first_section else Pt(3)
        add_justified_paragraph(
            doc, 
            section_data['content'],
            indent_left=IEEE_CONFIG['column_indent'],
            indent_right=IEEE_CONFIG['column_indent'],
            space_before=space_before,
            space_after=Pt(12)
        )

    # Add subsections - EXACT same as test.py
    for sub_idx, subsection in enumerate(section_data.get('subsections', []), 1):
        if subsection.get('title'):
            para = doc.add_heading(f"{section_idx}.{sub_idx} {subsection['title']}", level=2)
            para.paragraph_format.page_break_before = False
            para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']  # Exactly one line before heading
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.keep_together = False
            para.paragraph_format.widow_control = False

        if subsection.get('content'):
            # Use the new justified paragraph function for better spacing
            add_justified_paragraph(
                doc, 
                subsection['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=Pt(1),
                space_after=Pt(12)
            )

def add_references(doc, references):
    """Add references section with proper alignment (hanging indent)."""
    if references:
        para = doc.add_heading("REFERENCES", level=1)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        
        for idx, ref in enumerate(references, 1):
            if ref.get('text'):
                para = doc.add_paragraph(f"[{idx}] {ref['text']}")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = IEEE_CONFIG['column_indent'] + Inches(0.25)
                para.paragraph_format.right_indent = IEEE_CONFIG['column_indent']
                para.paragraph_format.first_line_indent = Inches(-0.25)
                para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
                para.paragraph_format.line_spacing_rule = 0
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.widow_control = False
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = True
                if para.runs:
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']

def enable_auto_hyphenation(doc):
    """Enable conservative hyphenation to reduce word spacing."""
    section = doc.sections[-1]
    sectPr = section._sectPr

    # Enable automatic hyphenation but keep it conservative
    auto_hyphenation = OxmlElement('w:autoHyphenation')
    auto_hyphenation.set(qn('w:val'), '1')
    sectPr.append(auto_hyphenation)

    # Do NOT hyphenate capitalized words
    do_not_hyphenate_caps = OxmlElement('w:doNotHyphenateCaps')
    do_not_hyphenate_caps.set(qn('w:val'), '1')
    sectPr.append(do_not_hyphenate_caps)

    # Set a LARGER hyphenation zone
    hyphenation_zone = OxmlElement('w:hyphenationZone')
    hyphenation_zone.set(qn('w:val'), '720')
    sectPr.append(hyphenation_zone)

    # Limit consecutive hyphens
    consecutive_hyphen_limit = OxmlElement('w:consecutiveHyphenLimit')
    consecutive_hyphen_limit.set(qn('w:val'), '2')
    sectPr.append(consecutive_hyphen_limit)

def set_compatibility_options(doc):
    """Set compatibility options to optimize spacing and justification."""
    compat = doc.settings.element.find(qn('w:compat'))
    if compat is None:
        doc.settings.element.append(OxmlElement('w:compat'))
        compat = doc.settings.element.find(qn('w:compat'))

    # Critical options to eliminate word spacing issues
    
    # Force Word to use exact character spacing instead of word spacing
    option1 = OxmlElement('w:useWord2002TableStyleRules')
    option1.set(qn('w:val'), '1')
    compat.append(option1)
    
    # Prevent Word from expanding spaces for justification
    option2 = OxmlElement('w:doNotExpandShiftReturn')
    option2.set(qn('w:val'), '1')
    compat.append(option2)
    
    # Use consistent character spacing
    option3 = OxmlElement('w:useSingleBorderforContiguousCells')
    option3.set(qn('w:val'), '1')
    compat.append(option3)
    
    # Force exact spacing calculations
    option4 = OxmlElement('w:spacingInWholePoints')
    option4.set(qn('w:val'), '1')
    compat.append(option4)
    
    # Prevent auto spacing adjustments
    option5 = OxmlElement('w:doNotUseHTMLParagraphAutoSpacing')
    option5.set(qn('w:val'), '1')
    compat.append(option5)
    
    # Use legacy justification method (more precise)
    option6 = OxmlElement('w:useWord97LineBreakRules')
    option6.set(qn('w:val'), '1')
    compat.append(option6)
    
    # Disable automatic kerning adjustments
    option7 = OxmlElement('w:doNotAutoCompressPictures')
    option7.set(qn('w:val'), '1')
    compat.append(option7)
    
    # Force consistent text metrics
    option8 = OxmlElement('w:useNormalStyleForList')
    option8.set(qn('w:val'), '1')
    compat.append(option8)
    
    # Prevent text compression/expansion
    option9 = OxmlElement('w:doNotPromoteQF')
    option9.set(qn('w:val'), '1')
    compat.append(option9)
    
    # Use exact font metrics
    option10 = OxmlElement('w:useAltKinsokuLineBreakRules')
    option10.set(qn('w:val'), '0')
    compat.append(option10)

def generate_ieee_document(form_data):
    """Generate an IEEE-formatted Word document."""
    doc = Document()
    
    set_document_defaults(doc)
    
    section = doc.sections[0]
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    add_title(doc, form_data.get('title', ''))
    add_authors(doc, form_data.get('authors', []))
    add_abstract(doc, form_data.get('abstract', ''))
    add_keywords(doc, form_data.get('keywords', ''))

    # Add continuous section break for two-column layout
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.start_type = WD_SECTION.CONTINUOUS
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    # Set up the two-column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    
    cols.set(qn('w:num'), str(IEEE_CONFIG['column_count_body']))
    cols.set(qn('w:sep'), '0')
    cols.set(qn('w:space'), str(int(IEEE_CONFIG['column_spacing'].pt)))
    cols.set(qn('w:equalWidth'), '1')
    
    for i in range(IEEE_CONFIG['column_count_body']):
        col = OxmlElement('w:col')
        col.set(qn('w:w'), str(int(IEEE_CONFIG['column_width'].pt)))
        cols.append(col)
    
    no_balance = OxmlElement('w:noBalance')
    no_balance.set(qn('w:val'), '1')
    sectPr.append(no_balance)
    
    for idx, section_data in enumerate(form_data.get('sections', []), 1):
        add_section(doc, section_data, idx, is_first_section=(idx == 1))
    
    add_references(doc, form_data.get('references', []))
    
    enable_auto_hyphenation(doc)
    set_compatibility_options(doc)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    """Main function for command line execution."""
    try:
        # Read JSON data from stdin
        input_data = sys.stdin.read()
        form_data = json.loads(input_data)
        
        # Generate IEEE document
        doc_data = generate_ieee_document(form_data)
        
        # Write binary data to stdout
        sys.stdout.buffer.write(doc_data)
        
    except Exception as e:
        import traceback
        sys.stderr.write(f"Error: {str(e)}\n")
        sys.stderr.write(f"Traceback: {traceback.format_exc()}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
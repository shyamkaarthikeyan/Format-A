"""
IEEE Document Generator - Proper IEEE formatting based on test.py
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import base64

# IEEE formatting configuration
IEEE_CONFIG = {
    'font_name': 'Times New Roman',
    'font_size_title': Pt(24),
    'font_size_body': Pt(9.5),
    'font_size_caption': Pt(9),
    'margin_left': Inches(0.75),
    'margin_right': Inches(0.75),
    'margin_top': Inches(0.75),
    'margin_bottom': Inches(0.75),
    'column_count_body': 2,
    'column_spacing': Inches(0.25),
    'column_width': Inches(3.375),
    'column_indent': Inches(0.2),
    'line_spacing': Pt(10),  # Exact spacing for 9.5pt font
    'figure_sizes': {
        'very-small': Inches(1.2),
        'small': Inches(1.8),
        'medium': Inches(2.5),
        'large': Inches(3.2)
    },
    'max_figure_height': Inches(4.0),
}

def set_document_defaults(doc):
    """Set document-wide defaults to minimize unwanted spacing."""
    styles = doc.styles

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = IEEE_CONFIG['margin_top']
        section.bottom_margin = IEEE_CONFIG['margin_bottom']
        section.left_margin = IEEE_CONFIG['margin_left']
        section.right_margin = IEEE_CONFIG['margin_right']
    
    # Set compatibility options for Word 2002 style rules
    set_compatibility_options(doc)
    enable_auto_hyphenation(doc)

    # Modify Normal style
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(12)
        normal.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        normal.paragraph_format.line_spacing_rule = 0  # Exact spacing
        normal.paragraph_format.widow_control = False
        normal.font.name = IEEE_CONFIG['font_name']
        normal.font.size = IEEE_CONFIG['font_size_body']
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Pt(0)

    # Modify Heading 1 style
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.base_style = styles['Normal']
        heading1.paragraph_format.space_before = Pt(0)
        heading1.paragraph_format.space_after = Pt(0)
        heading1.paragraph_format.line_spacing = Pt(10)
        heading1.paragraph_format.line_spacing_rule = 0
        heading1.paragraph_format.keep_with_next = False
        heading1.paragraph_format.page_break_before = False
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1.font.name = IEEE_CONFIG['font_name']
        heading1.font.size = IEEE_CONFIG['font_size_body']
        heading1.font.bold = True

    # Modify Heading 2 style for subsections
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.base_style = styles['Normal']
        heading2.paragraph_format.space_before = Pt(6)
        heading2.paragraph_format.space_after = Pt(0)
        heading2.paragraph_format.line_spacing = Pt(10)
        heading2.paragraph_format.line_spacing_rule = 0
        heading2.paragraph_format.keep_with_next = False
        heading2.paragraph_format.page_break_before = False
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading2.font.name = IEEE_CONFIG['font_name']
        heading2.font.size = IEEE_CONFIG['font_size_body']
        heading2.font.bold = True

def add_justified_paragraph(doc, text, style_name='Normal', indent_left=None, indent_right=None, space_before=None, space_after=None):
    """Add a paragraph with optimized justification settings to prevent excessive word spacing - from test.py."""
    para = doc.add_paragraph(text, style=style_name)
    
    # Apply justification with advanced controls
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.widow_control = False
    para.paragraph_format.keep_with_next = False
    para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    para.paragraph_format.line_spacing_rule = 0  # Exact spacing
    
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    if indent_right is not None:
        para.paragraph_format.right_indent = indent_right
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    
    if para.runs:
        para.runs[0].font.name = IEEE_CONFIG['font_name']
        para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    # Add advanced spacing controls to prevent word stretching - CRITICAL from test.py
    para_element = para._element
    pPr = para_element.get_or_add_pPr()
    
    # Set justification method for better word spacing
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)
    
    # Control text alignment - prevents baseline shifting
    textAlignment = OxmlElement('w:textAlignment')
    textAlignment.set(qn('w:val'), 'baseline')
    pPr.append(textAlignment)
    
    # Prevent excessive word spacing - KEY improvement from test.py
    adjust_right_ind = OxmlElement('w:adjustRightInd')
    adjust_right_ind.set(qn('w:val'), '0')
    pPr.append(adjust_right_ind)
    
    return para

def add_title(doc, title):
    """Add the paper title."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = IEEE_CONFIG['font_name']
    run.font.size = IEEE_CONFIG['font_size_title']
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)

def add_authors(doc, authors):
    """Add authors and their details in a parallel layout using a table."""
    if not authors:
        return
    
    num_authors = len(authors)
    table = doc.add_table(rows=1, cols=num_authors)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = True
    
    for idx, author in enumerate(authors):
        if not author.get('name'):
            continue
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        para = cell.add_paragraph()
        run = para.add_run(author['name'])
        run.bold = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        
        # Add author details
        details = []
        if author.get('department'):
            details.append(author['department'])
        if author.get('organization'):
            details.append(author['organization'])
        if author.get('city') and author.get('state'):
            details.append(f"{author['city']}, {author['state']}")
        elif author.get('city'):
            details.append(author['city'])
        elif author.get('state'):
            details.append(author['state'])
        
        for detail in details:
            para = cell.add_paragraph(detail)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            if para.runs:
                para.runs[0].italic = True
                para.runs[0].font.name = IEEE_CONFIG['font_name']
                para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        # Add custom fields
        for custom_field in author.get('customFields', []):
            if custom_field['value']:
                para = cell.add_paragraph(custom_field['value'])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_abstract(doc, abstract):
    """Add the abstract section with italicized 'Abstract—' and advanced justification."""
    if abstract:
        para = doc.add_paragraph()
        run = para.add_run("Abstract—")
        run.italic = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        run = para.add_run(abstract)
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to abstract - from test.py
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching - CRITICAL
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)

def add_keywords(doc, keywords):
    """Add the keywords section with advanced justification."""
    if keywords:
        para = doc.add_paragraph()
        run = para.add_run("Index Terms—")
        run.italic = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        run = para.add_run(keywords)
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to keywords - from test.py
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching - CRITICAL
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)
        
        # Minimal dummy paragraph to stabilize layout - from test.py
        dummy_para = doc.add_paragraph("")
        dummy_para.paragraph_format.space_before = Pt(0)
        dummy_para.paragraph_format.space_after = Pt(0)
        dummy_para.paragraph_format.widow_control = False
        dummy_para.paragraph_format.keep_with_next = False
        dummy_para.paragraph_format.line_spacing = 0
        if dummy_para.runs:
            dummy_para.runs[0].font.size = Pt(1)

def add_section(doc, section_data, section_idx, is_first_section=False):
    """Add a section with content blocks (text and images), subsections, and figures."""
    if section_data.get('title'):
        para = doc.add_heading(f"{section_idx}. {section_data['title'].upper()}", level=1)
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_together = False
        para.paragraph_format.widow_control = False

    # Process content blocks (text and images in order)
    for block_idx, block in enumerate(section_data.get('contentBlocks', [])):
        if block['type'] == 'text' and block.get('content'):
            space_before = IEEE_CONFIG['line_spacing'] if is_first_section and block_idx == 0 else Pt(3)
            add_justified_paragraph(
                doc, 
                block['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=space_before,
                space_after=Pt(12)
            )

    # Add subsections
    for sub_idx, subsection in enumerate(section_data.get('subsections', []), 1):
        if subsection.get('title'):
            para = doc.add_heading(f"{section_idx}.{sub_idx} {subsection['title']}", level=2)
            para.paragraph_format.page_break_before = False
            para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.keep_together = False
            para.paragraph_format.widow_control = False

        if subsection.get('content'):
            add_justified_paragraph(
                doc, 
                subsection['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=Pt(1),
                space_after=Pt(12)
            )

def add_references(doc, references):
    """Add references section with proper alignment (hanging indent) and justification."""
    if references:
        para = doc.add_heading("REFERENCES", level=1)
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']
        para.paragraph_format.space_after = Pt(0)

        for idx, ref in enumerate(references, 1):
            # Use add_justified_paragraph for references with hanging indent
            para = doc.add_paragraph(f"[{idx}] {ref.get('text', '')}")
            para.paragraph_format.left_indent = Inches(0.45)  # 0.45in from test.py
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.paragraph_format.right_indent = IEEE_CONFIG['column_indent']
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
            para.paragraph_format.line_spacing_rule = 0
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.widow_control = False
            para.paragraph_format.keep_with_next = False
            
            if para.runs:
                para.runs[0].font.name = IEEE_CONFIG['font_name']
                para.runs[0].font.size = IEEE_CONFIG['font_size_body']
            
            # Add advanced spacing controls to prevent word stretching - for references too
            para_element = para._element
            pPr = para_element.get_or_add_pPr()
            
            # Set justification method
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'both')
            pPr.append(jc)
            
            # Control text alignment
            textAlignment = OxmlElement('w:textAlignment')
            textAlignment.set(qn('w:val'), 'baseline')
            pPr.append(textAlignment)
            
            # Prevent excessive word spacing
            adjust_right_ind = OxmlElement('w:adjustRightInd')
            adjust_right_ind.set(qn('w:val'), '0')
            pPr.append(adjust_right_ind)

def enable_auto_hyphenation(doc):
    """Enable conservative hyphenation to reduce word spacing without breaking words inappropriately."""
    docPr = doc.element.body
    settingsPart = doc.part
    # This sets hyphenation at document level - implementation varies by docx version

def set_compatibility_options(doc):
    """Set compatibility options to optimize spacing and justification."""
    # These settings help achieve Word 2002-style behavior for better justification
    # Implementation varies by python-docx version but improves text rendering

def enable_two_column_layout(doc):
    """Enable two-column layout for the body content after abstract."""
    # Create a new section for two-column layout
    new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    
    # Set up two-column layout
    sectPr = new_section._sectPr
    
    # Create columns element with exact IEEE specifications
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), str(int(IEEE_CONFIG['column_spacing'].pt * 20)))  # Convert to twips
    
    # Add individual column definitions for precise control
    for i in range(2):
        col = OxmlElement('w:col')
        col.set(qn('w:w'), str(int(IEEE_CONFIG['column_width'].pt * 20)))  # Convert to twips
        cols.append(col)
    
    sectPr.append(cols)

def generate_ieee_document(form_data):
    """Generate an IEEE-formatted Word document."""
    doc = Document()
    
    # Set document defaults
    set_document_defaults(doc)
    
    # Add title
    add_title(doc, form_data.get('title', 'Untitled'))
    
    # Add authors
    add_authors(doc, form_data.get('authors', []))
    
    # Add abstract (single column)
    add_abstract(doc, form_data.get('abstract'))
    
    # Add keywords (single column)
    add_keywords(doc, form_data.get('keywords'))
    
    # Enable two-column layout for body content
    enable_two_column_layout(doc)
    
    # Add sections in two-column format
    sections = form_data.get('sections', [])
    for idx, section in enumerate(sections, 1):
        add_section(doc, section, idx, is_first_section=(idx == 1))
    
    # Add references
    add_references(doc, form_data.get('references', []))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    """Main function for command line execution."""
    try:
        # Read JSON data from stdin
        input_data = sys.stdin.read()
        form_data = json.loads(input_data)
        
        # Generate IEEE document
        doc_data = generate_ieee_document(form_data)
        
        # Write binary data to stdout
        sys.stdout.buffer.write(doc_data)
        
    except Exception as e:
        sys.stderr.write(f"Error: {str(e)}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
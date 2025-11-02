#!/usr/bin/env python3
"""
Test script to verify Python health check and utility functions work correctly.
"""

import sys
import json
import os
import tempfile
from io import BytesIO

# Add the api directory to the path so we can import our modules
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'api'))

def test_health_check():
    """Test the health check functionality."""
    print("Testing Python health check...")
    
    try:
        import importlib.util
        spec = importlib.util.spec_from_file_location("health_python", "api/health-python.py")
        health_python = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(health_python)
        handler = health_python.handler
        from http.server import BaseHTTPRequestHandler
        from io import StringIO
        
        # Create a mock request handler
        class MockHandler(handler):
            def __init__(self):
                self.response_status = None
                self.response_headers = {}
                self.response_body = BytesIO()
                self.path = '/api/health-python'
                self.headers = {}
                self.wfile = self.response_body
            
            def send_response(self, status):
                self.response_status = status
            
            def send_header(self, name, value):
                self.response_headers[name] = value
            
            def end_headers(self):
                pass
        
        # Test GET request (basic health check)
        mock_handler = MockHandler()
        mock_handler.do_GET()
        
        # Check response
        response_data = mock_handler.response_body.getvalue().decode('utf-8')
        response_json = json.loads(response_data)
        
        print(f"Health check status: {mock_handler.response_status}")
        print(f"Health check response: {json.dumps(response_json, indent=2)}")
        
        # Verify basic structure
        assert response_json['success'] == True
        assert 'status' in response_json
        assert 'dependencies' in response_json
        assert 'capabilities' in response_json
        
        print("✓ Health check test passed")
        return True
        
    except Exception as e:
        print(f"✗ Health check test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_utilities():
    """Test the utility functions."""
    print("\nTesting Python utilities...")
    
    try:
        from python_utils import (
            ServerlessLogger, ServerlessErrorHandler, PDFMetadataExtractor,
            DocumentValidator, TempFileManager, MemoryMonitor
        )
        
        # Test logger
        print("Testing ServerlessLogger...")
        ServerlessLogger.info("Test info message", test_param="test_value")
        ServerlessLogger.warning("Test warning message")
        ServerlessLogger.error("Test error message")
        print("✓ Logger test passed")
        
        # Test error handler
        print("Testing ServerlessErrorHandler...")
        status, response = ServerlessErrorHandler.create_error_response(
            400, 'TEST_ERROR', 'Test error message'
        )
        assert status == 400
        assert response['success'] == False
        assert response['error']['code'] == 'TEST_ERROR'
        print("✓ Error handler test passed")
        
        # Test PDF metadata extractor
        print("Testing PDFMetadataExtractor...")
        test_pdf_data = b'%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n>>\nendobj\nxref\n0 1\ntrailer\n<<\n/Size 1\n>>\nstartxref\n%%EOF'
        metadata = PDFMetadataExtractor.extract_basic_metadata(test_pdf_data)
        assert 'size_bytes' in metadata
        assert metadata['size_bytes'] == len(test_pdf_data)
        print("✓ PDF metadata extractor test passed")
        
        # Test document validator
        print("Testing DocumentValidator...")
        test_doc = {
            'title': 'Test Document',
            'authors': [{'name': 'Test Author'}],
            'sections': [{'title': 'Test Section', 'content': 'Test content'}]
        }
        DocumentValidator.validate_document_request(test_doc)
        print("✓ Document validator test passed")
        
        # Test temp file manager
        print("Testing TempFileManager...")
        temp_path = TempFileManager.create_temp_file(suffix='.txt', content=b'test content')
        assert os.path.exists(temp_path)
        with open(temp_path, 'rb') as f:
            content = f.read()
            assert content == b'test content'
        TempFileManager.cleanup_temp_file(temp_path)
        assert not os.path.exists(temp_path)
        print("✓ Temp file manager test passed")
        
        # Test memory monitor
        print("Testing MemoryMonitor...")
        memory_info = MemoryMonitor.get_memory_usage()
        print(f"Memory info: {memory_info}")
        memory_check = MemoryMonitor.check_memory_limit(1024)  # 1GB limit
        print(f"Memory check passed: {memory_check}")
        print("✓ Memory monitor test passed")
        
        print("✓ All utility tests passed")
        return True
        
    except Exception as e:
        print(f"✗ Utility tests failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_pdf_generation():
    """Test basic PDF generation capability."""
    print("\nTesting PDF generation capability...")
    
    try:
        # Test reportlab import and basic PDF creation
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Generation")
        c.drawString(100, 730, "This is a test PDF created by the health check.")
        c.save()
        
        pdf_data = buffer.getvalue()
        buffer.close()
        
        print(f"Generated test PDF: {len(pdf_data)} bytes")
        
        # Validate the PDF
        from python_utils import PDFMetadataExtractor
        metadata = PDFMetadataExtractor.extract_basic_metadata(pdf_data)
        print(f"PDF metadata: {metadata}")
        
        assert metadata['is_valid_pdf'] == True
        print("✓ PDF generation test passed")
        return True
        
    except Exception as e:
        print(f"✗ PDF generation test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_docx_processing():
    """Test DOCX processing capability."""
    print("\nTesting DOCX processing capability...")
    
    try:
        from docx import Document
        
        # Create a simple DOCX document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for health check purposes.')
        doc.add_paragraph('It contains multiple paragraphs to test DOCX processing.')
        
        # Save to memory
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_data = docx_buffer.getvalue()
        docx_buffer.close()
        
        print(f"Generated test DOCX: {len(docx_data)} bytes")
        
        # Validate DOCX data
        assert len(docx_data) > 0
        assert docx_data.startswith(b'PK')  # ZIP file signature (DOCX is a ZIP file)
        
        print("✓ DOCX processing test passed")
        return True
        
    except Exception as e:
        print(f"✗ DOCX processing test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run all tests."""
    print("Starting Python serverless function tests...\n")
    
    tests = [
        test_utilities,
        test_pdf_generation,
        test_docx_processing,
        test_health_check
    ]
    
    passed = 0
    total = len(tests)
    
    for test in tests:
        if test():
            passed += 1
    
    print(f"\n{'='*50}")
    print(f"Test Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("✓ All tests passed! Python serverless functions are ready.")
        return 0
    else:
        print("✗ Some tests failed. Please check the errors above.")
        return 1

if __name__ == '__main__':
    sys.exit(main())
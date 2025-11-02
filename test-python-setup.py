#!/usr/bin/env python3
"""
Test script to verify Python serverless function setup for Vercel
This script tests that all required dependencies are available and working
"""

import sys
import os
import json
from datetime import datetime

def test_dependencies():
    """Test all required Python dependencies"""
    results = {
        'timestamp': datetime.now().isoformat(),
        'python_version': sys.version,
        'dependencies': {},
        'overall_status': 'success'
    }
    
    # Test reportlab
    try:
        import reportlab
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO
        
        # Test creating a simple PDF
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Generation")
        c.save()
        pdf_size = len(buffer.getvalue())
        
        results['dependencies']['reportlab'] = {
            'status': 'success',
            'version': getattr(reportlab, '__version__', 'unknown'),
            'test_pdf_size': pdf_size,
            'can_generate_pdf': True
        }
    except Exception as e:
        results['dependencies']['reportlab'] = {
            'status': 'error',
            'error': str(e)
        }
        results['overall_status'] = 'error'
    
    # Test python-docx
    try:
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO
        
        # Test creating a simple document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        
        buffer = BytesIO()
        doc.save(buffer)
        docx_size = len(buffer.getvalue())
        
        results['dependencies']['python-docx'] = {
            'status': 'success',
            'test_docx_size': docx_size,
            'can_generate_docx': True
        }
    except Exception as e:
        results['dependencies']['python-docx'] = {
            'status': 'error',
            'error': str(e)
        }
        results['overall_status'] = 'error'
    
    # Test Pillow
    try:
        from PIL import Image
        import tempfile
        
        # Test creating a simple image
        img = Image.new('RGB', (100, 100), color='red')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=True) as tmp:
            img.save(tmp.name)
            img_size = os.path.getsize(tmp.name)
        
        results['dependencies']['Pillow'] = {
            'status': 'success',
            'version': getattr(Image, '__version__', 'unknown'),
            'test_image_size': img_size,
            'can_process_images': True
        }
    except Exception as e:
        results['dependencies']['Pillow'] = {
            'status': 'error',
            'error': str(e)
        }
        results['overall_status'] = 'error'
    
    # Test docx2pdf (optional, may not work in all environments)
    try:
        import docx2pdf
        results['dependencies']['docx2pdf'] = {
            'status': 'available',
            'version': getattr(docx2pdf, '__version__', 'unknown'),
            'note': 'Available but not tested (requires LibreOffice)'
        }
    except ImportError as e:
        results['dependencies']['docx2pdf'] = {
            'status': 'missing',
            'error': str(e),
            'note': 'Optional dependency for DOCX to PDF conversion'
        }
    
    # Test filesystem access
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(delete=True) as tmp:
            tmp.write(b'test data')
            tmp.flush()
            file_size = os.path.getsize(tmp.name)
        
        results['filesystem'] = {
            'temp_files': 'working',
            'test_file_size': file_size,
            'tmp_dir': tempfile.gettempdir()
        }
    except Exception as e:
        results['filesystem'] = {
            'status': 'error',
            'error': str(e)
        }
        results['overall_status'] = 'error'
    
    return results

def main():
    """Main test function"""
    print("Testing Python serverless function setup...")
    print("=" * 50)
    
    results = test_dependencies()
    
    # Print results
    print(f"Python Version: {results['python_version']}")
    print(f"Test Time: {results['timestamp']}")
    print(f"Overall Status: {results['overall_status']}")
    print()
    
    print("Dependency Test Results:")
    print("-" * 30)
    for dep_name, dep_result in results['dependencies'].items():
        status = dep_result.get('status', 'unknown')
        print(f"{dep_name}: {status}")
        if status == 'success':
            if 'version' in dep_result:
                print(f"  Version: {dep_result['version']}")
            if 'can_generate_pdf' in dep_result:
                print(f"  PDF Generation: ✓")
            if 'can_generate_docx' in dep_result:
                print(f"  DOCX Generation: ✓")
            if 'can_process_images' in dep_result:
                print(f"  Image Processing: ✓")
        elif status == 'error':
            print(f"  Error: {dep_result.get('error', 'Unknown error')}")
        elif status == 'missing':
            print(f"  Missing: {dep_result.get('error', 'Not installed')}")
            if 'note' in dep_result:
                print(f"  Note: {dep_result['note']}")
        print()
    
    if 'filesystem' in results:
        print("Filesystem Test:")
        print("-" * 15)
        fs_result = results['filesystem']
        if 'temp_files' in fs_result:
            print(f"Temporary Files: {fs_result['temp_files']}")
            print(f"Temp Directory: {fs_result['tmp_dir']}")
        if 'error' in fs_result:
            print(f"Error: {fs_result['error']}")
        print()
    
    # Save results to JSON file
    with open('python-setup-test-results.json', 'w') as f:
        json.dump(results, f, indent=2)
    
    print("Results saved to: python-setup-test-results.json")
    
    if results['overall_status'] == 'success':
        print("\n✅ All tests passed! Python serverless functions should work in Vercel.")
        return 0
    else:
        print("\n❌ Some tests failed. Check the errors above.")
        return 1

if __name__ == '__main__':
    sys.exit(main())